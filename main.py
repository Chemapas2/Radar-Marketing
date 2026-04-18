import io
import os
import re
from collections import Counter
from datetime import date, datetime, timezone
from typing import Dict, List, Optional, Sequence
from urllib.parse import parse_qs, quote, unquote, urlparse

import feedparser
import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup
from dateutil import parser as date_parser
from docx import Document

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None


APP_TITLE = "Nutreco Iberia | Radar de necesidades del mercado"
USER_AGENT = "Mozilla/5.0 (compatible; NutrecoRadar/2.0; +https://streamlit.io)"
REQUEST_TIMEOUT = 20
DEFAULT_COMPANY_CONTEXT = """Objetivo del radar:
- Leer necesidades del mercado para orientar productos, servicios y soluciones de Nutreco Iberia.
- Detectar señales comerciales, técnicas y regulatorias que puedan traducirse en innovación, soporte técnico o argumentario de valor.
- Priorizar oportunidades accionables y riesgos que afecten a posicionamiento, portafolio o claims.
"""

MARKET_NEWS_SITES = [
    "mapa.gob.es",
    "agroinformacion.com",
    "eurocarne.com",
    "diarioveterinario.com",
    "interempresas.net",
    "agrodigital.com",
    "animal-health-media.com",
]
OFFICIAL_REG_DOMAINS = [
    "boe.es",
    "eur-lex.europa.eu",
    "mapa.gob.es",
    "aesan.gob.es",
    "efsa.europa.eu",
    "miteco.gob.es",
    "sanidad.gob.es",
    "europa.eu",
]
SCIENCE_HINTS = {
    "study",
    "review",
    "trial",
    "journal",
    "effect",
    "effects",
    "performance",
    "animals",
    "animal",
    "feed",
    "nutrition",
    "dairy",
    "swine",
    "pig",
    "broiler",
    "laying",
    "ruminant",
    "sheep",
    "goat",
    "cattle",
    "rabbit",
}
STOPWORDS = {
    "de", "la", "el", "los", "las", "y", "o", "u", "en", "por", "para", "con", "sin", "del", "al", "a",
    "the", "and", "for", "with", "from", "into", "this", "that", "animal", "animals", "market", "news",
    "regulation", "science", "technical", "sector", "ganaderia", "ganadería", "livestock", "production",
}

SPECIES_PROFILES: Dict[str, Dict[str, List[str]]] = {
    "All species": {
        "aliases": ["ganadería", "livestock", "animal production", "farm animals", "ruminants", "poultry", "swine", "rabbit"],
        "market_aliases": ["ganadería", "livestock", "alimentación animal", "animal nutrition", "feed sector"],
        "science_aliases": ["livestock", "animal nutrition", "animal production", "farm animals"],
        "reg_aliases": ["ganadería", "alimentación animal", "livestock", "animal health", "animal welfare"],
    },
    "Alimentación animal": {
        "aliases": ["alimentación animal", "animal nutrition", "feed", "compound feed", "feed industry", "piensos"],
        "market_aliases": ["alimentación animal", "piensos", "feed sector", "feed industry", "compound feed"],
        "science_aliases": ["animal nutrition", "feed", "feed additives", "compound feed", "animal feeding"],
        "reg_aliases": ["feed", "piensos", "alimentación animal", "feed hygiene", "feed additives", "compound feed"],
    },
    "Avicultura de puesta": {
        "aliases": ["avicultura de puesta", "gallinas ponedoras", "layers", "laying hens", "egg production"],
        "market_aliases": ["gallinas ponedoras", "layers", "huevo", "egg market", "egg production"],
        "science_aliases": ["laying hens", "layers", "egg production", "egg quality"],
        "reg_aliases": ["laying hens", "gallinas ponedoras", "huevo", "egg marketing"],
    },
    "Avicultura de carne": {
        "aliases": ["avicultura de carne", "pollos de engorde", "broilers", "broiler production", "meat poultry"],
        "market_aliases": ["broilers", "pollo", "broiler market", "pollos de engorde"],
        "science_aliases": ["broilers", "broiler production", "pollo de engorde", "broiler nutrition"],
        "reg_aliases": ["broilers", "pollos de engorde", "avicultura de carne", "broiler welfare"],
    },
    "Porcino": {
        "aliases": ["porcino", "swine", "pig", "pig production", "cerdo"],
        "market_aliases": ["porcino", "cerdo", "swine market", "pig market"],
        "science_aliases": ["swine", "pig", "piglets", "sows", "porcine"],
        "reg_aliases": ["porcino", "cerdo", "swine", "porcine"],
    },
    "Vacuno de leche": {
        "aliases": ["vacuno de leche", "dairy cattle", "dairy cows", "lechero", "milk production"],
        "market_aliases": ["vacuno de leche", "leche", "dairy market", "dairy cows"],
        "science_aliases": ["dairy cows", "dairy cattle", "transition cows", "milk yield"],
        "reg_aliases": ["vacuno de leche", "dairy cows", "milk production", "dairy sector"],
    },
    "Vacuno de carne": {
        "aliases": ["vacuno de carne", "beef cattle", "beef production", "cebo", "beef"],
        "market_aliases": ["vacuno de carne", "beef cattle", "beef market", "beef"],
        "science_aliases": ["beef cattle", "beef production", "feedlot", "growing-finishing cattle"],
        "reg_aliases": ["vacuno de carne", "beef cattle", "beef", "bovine"],
    },
    "Ovino": {
        "aliases": ["ovino", "sheep", "lamb", "sheep production", "cordero"],
        "market_aliases": ["ovino", "cordero", "sheep market", "lamb market"],
        "science_aliases": ["sheep", "lamb", "ewes", "small ruminants"],
        "reg_aliases": ["ovino", "sheep", "lamb", "small ruminants"],
    },
    "Caprino": {
        "aliases": ["caprino", "goat", "goats", "goat production", "leche caprina"],
        "market_aliases": ["caprino", "goat", "goat market", "leche caprina"],
        "science_aliases": ["goat", "goats", "dairy goats", "small ruminants"],
        "reg_aliases": ["caprino", "goat", "goats", "small ruminants"],
    },
    "Cunicultura": {
        "aliases": ["cunicultura", "rabbit", "rabbits", "rabbit production", "conejo"],
        "market_aliases": ["conejo", "rabbit market", "cunicultura", "rabbit production"],
        "science_aliases": ["rabbit", "rabbits", "cuniculture", "rabbit nutrition"],
        "reg_aliases": ["cunicultura", "rabbit", "rabbits", "conejo"],
    },
}

CATEGORY_LABELS = {
    "market": "Mercado y noticias",
    "science": "Científico-técnico",
    "regulation": "Legislación y regulación",
}

MARKET_TOPIC_ITEMS = [
    {"label": "Precios y cotizaciones", "terms": ["precios", "cotizaciones", "price", "market prices"]},
    {"label": "Boletines oficiales de precios", "terms": ["boletín de precios", "precios semanales", "cotizaciones oficiales", "mercados agrarios"]},
    {"label": "Coste del pienso", "terms": ["coste del pienso", "feed cost", "piensos", "compound feed prices"]},
    {"label": "Materias primas y piensos", "terms": ["materias primas", "piensos", "feed materials", "compound feed"]},
    {"label": "Cereales y energía", "terms": ["cereales", "maíz", "trigo", "barley", "grain prices"]},
    {"label": "Harina de soja y proteínas", "terms": ["soja", "soybean meal", "protein meals", "oleaginosas"]},
    {"label": "Aceites y grasas", "terms": ["aceites", "grasas", "fats and oils", "oleínas"]},
    {"label": "Aditivos y especialidades", "terms": ["aditivos", "especialidades", "feed additives", "specialties"]},
    {"label": "Rentabilidad y márgenes", "terms": ["rentabilidad", "márgenes", "profitability", "margins"]},
    {"label": "Coste energético", "terms": ["energía", "electricidad", "gas", "energy costs"]},
    {"label": "Inflación alimentaria", "terms": ["inflación alimentaria", "food inflation", "consumer prices"]},
    {"label": "Consumo interior", "terms": ["consumo", "demanda interna", "domestic demand", "consumption"]},
    {"label": "Tendencias del consumidor", "terms": ["consumer trends", "tendencias del consumidor", "shopper", "market demand"]},
    {"label": "Retail y gran distribución", "terms": ["retail", "gran distribución", "supermercados", "modern trade"]},
    {"label": "Canal horeca y foodservice", "terms": ["horeca", "foodservice", "restauración", "catering"]},
    {"label": "Exportación", "terms": ["exportación", "exports", "overseas sales", "trade"]},
    {"label": "Importación", "terms": ["importación", "imports", "trade flows", "sourcing"]},
    {"label": "Comercio internacional", "terms": ["comercio internacional", "international trade", "global market", "trade flows"]},
    {"label": "Aranceles y barreras comerciales", "terms": ["aranceles", "trade barriers", "customs", "restricciones comerciales"]},
    {"label": "Logística y transporte", "terms": ["logística", "transporte", "freight", "shipping"]},
    {"label": "Sacrificio y mataderos", "terms": ["mataderos", "sacrificio", "slaughter", "abattoirs"]},
    {"label": "Producción sectorial", "terms": ["producción", "output", "sector production", "supply"]},
    {"label": "Censo y capacidad productiva", "terms": ["censo", "inventario ganadero", "herd size", "flock size"]},
    {"label": "Disponibilidad de animales", "terms": ["disponibilidad de animales", "placements", "animal supply", "reposiciones"]},
    {"label": "Noticias del MAPA y del sector", "terms": ["MAPA", "ministerio de agricultura", "sector news", "ganadería noticias"]},
    {"label": "Coyuntura agraria", "terms": ["coyuntura", "agrarian outlook", "market outlook", "situación de mercado"]},
    {"label": "Boletines lácteos", "terms": ["boletín lácteo", "precio leche", "dairy bulletin", "milk market"]},
    {"label": "Boletines cárnicos", "terms": ["boletín cárnico", "carne", "meat market", "meat bulletin"]},
    {"label": "Precio de la leche", "terms": ["precio leche", "milk price", "farmgate milk", "dairy prices"]},
    {"label": "Precio del cerdo", "terms": ["precio cerdo", "pig price", "hog prices", "porcino precios"]},
    {"label": "Precio del pollo", "terms": ["precio pollo", "broiler price", "poultry prices", "pollo mercado"]},
    {"label": "Precio del huevo", "terms": ["precio huevo", "egg price", "egg market", "cotización huevo"]},
    {"label": "Precio del vacuno", "terms": ["precio vacuno", "beef price", "cattle prices", "vacuno mercado"]},
    {"label": "Precio del cordero", "terms": ["precio cordero", "lamb price", "ovine market", "cordero mercado"]},
    {"label": "Precio del conejo", "terms": ["precio conejo", "rabbit price", "cunicultura mercado", "conejo mercado"]},
    {"label": "Precio de leche de cabra", "terms": ["precio leche cabra", "goat milk price", "caprino mercado", "leche caprina precio"]},
    {"label": "Calidad percibida y valor", "terms": ["calidad", "value perception", "consumer value", "premium quality"]},
    {"label": "Bienestar animal como driver de mercado", "terms": ["bienestar animal", "welfare demand", "consumer welfare", "animal welfare market"]},
    {"label": "Sostenibilidad como driver de compra", "terms": ["sostenibilidad", "sustainable demand", "low carbon", "green demand"]},
    {"label": "Huella de carbono de producto", "terms": ["huella de carbono", "carbon footprint", "scope 3", "low carbon"]},
    {"label": "Certificaciones y sellos", "terms": ["certificaciones", "sellos", "certification", "quality label"]},
    {"label": "Producto premium", "terms": ["premium", "high value", "specialty", "valor añadido"]},
    {"label": "Segmento económico y commodity", "terms": ["commodity", "económico", "value segment", "precio bajo"]},
    {"label": "Bioseguridad con impacto de mercado", "terms": ["bioseguridad", "disease impact market", "trade restrictions", "market impact"]},
    {"label": "Enfermedades y restricciones comerciales", "terms": ["restricciones comerciales", "sanidad animal", "disease outbreak", "trade ban"]},
    {"label": "Resistencia del consumidor a antibióticos", "terms": ["antibióticos", "antibiotic free", "raised without antibiotics", "consumer concern"]},
    {"label": "Innovación de producto", "terms": ["innovación", "product innovation", "new solutions", "launch"]},
    {"label": "Ingredientes funcionales", "terms": ["ingredientes funcionales", "functional ingredients", "specialty feed", "health solutions"]},
    {"label": "Proteínas alternativas", "terms": ["proteínas alternativas", "alternative proteins", "insect meal", "novel feeds"]},
    {"label": "Digitalización y datos en granja", "terms": ["digitalización", "data", "precision farming", "sensors"]},
    {"label": "Financiación e inversión", "terms": ["inversión", "financiación", "capex", "farm investment"]},
    {"label": "Fusiones y adquisiciones", "terms": ["adquisiciones", "fusiones", "mergers", "acquisition"]},
    {"label": "Empleo y relevo generacional", "terms": ["empleo", "relevo generacional", "labour", "farm succession"]},
    {"label": "Seguros y gestión del riesgo", "terms": ["seguros", "risk management", "insurance", "volatility"]},
    {"label": "Disponibilidad de agua", "terms": ["agua", "water availability", "drought", "water stress"]},
    {"label": "Clima y eventos extremos", "terms": ["clima", "heat waves", "drought", "extreme weather"]},
    {"label": "Compras públicas y comedor colectivo", "terms": ["compras públicas", "public procurement", "school meals", "institutional demand"]},
]

SCIENCE_TOPIC_ITEMS = [
    {"label": "Nutrición y formulación", "terms": ["nutrition", "nutrición", "formulation", "diet"]},
    {"label": "Eficiencia alimentaria e índice de conversión", "terms": ["feed efficiency", "feed conversion", "índice de conversión", "FCR"]},
    {"label": "Consumo voluntario e ingestión", "terms": ["feed intake", "ingestión", "voluntary intake", "appetite"]},
    {"label": "Salud intestinal", "terms": ["gut health", "salud intestinal", "intestinal integrity", "microbiota"]},
    {"label": "Microbiota y moduladores", "terms": ["microbiota", "microbiome", "microbial modulation", "intestinal microbiota"]},
    {"label": "Probióticos", "terms": ["probiotics", "probióticos", "direct fed microbials", "beneficial bacteria"]},
    {"label": "Prebióticos", "terms": ["prebiotics", "prebióticos", "oligosaccharides", "mannan"]},
    {"label": "Simbióticos", "terms": ["synbiotics", "simbióticos", "combined probiotic prebiotic", "microbiota support"]},
    {"label": "Ácidos orgánicos", "terms": ["organic acids", "ácidos orgánicos", "acidifiers", "acidification"]},
    {"label": "Fitogénicos y aceites esenciales", "terms": ["phytogenics", "fitogénicos", "essential oils", "plant extracts"]},
    {"label": "Enzimas", "terms": ["enzymes", "enzimas", "phytase", "xylanase"]},
    {"label": "Minerales y oligoelementos", "terms": ["minerals", "trace minerals", "oligoelementos", "microminerals"]},
    {"label": "Vitaminas", "terms": ["vitamins", "vitaminas", "fat-soluble vitamins", "supplementation"]},
    {"label": "Aminoácidos", "terms": ["amino acids", "aminoácidos", "lysine", "methionine"]},
    {"label": "Proteína y aminoácidos ideales", "terms": ["ideal protein", "protein reduction", "amino acid balance", "low protein diet"]},
    {"label": "Fibra y estructura de la dieta", "terms": ["fiber", "fibra", "diet structure", "physically effective fiber"]},
    {"label": "Grasa y energía neta", "terms": ["fat supplementation", "energy density", "net energy", "dietary fat"]},
    {"label": "Micotoxinas", "terms": ["mycotoxins", "micotoxinas", "aflatoxin", "deoxynivalenol"]},
    {"label": "Calidad del agua", "terms": ["water quality", "calidad del agua", "drinking water", "water sanitation"]},
    {"label": "Bioseguridad y prevención", "terms": ["biosecurity", "bioseguridad", "prevention", "farm hygiene"]},
    {"label": "Reducción de antimicrobianos", "terms": ["antimicrobial reduction", "uso prudente", "antibiotic reduction", "AMR"]},
    {"label": "Vacunación e inmunidad", "terms": ["vaccination", "vacunación", "immunity", "immune response"]},
    {"label": "Estrés térmico", "terms": ["heat stress", "estrés térmico", "thermal stress", "cooling"]},
    {"label": "Bienestar animal", "terms": ["animal welfare", "bienestar animal", "behavior", "housing"]},
    {"label": "Sostenibilidad y eficiencia ambiental", "terms": ["sustainability", "environmental efficiency", "life cycle", "LCA"]},
    {"label": "Metano y huella de carbono", "terms": ["methane", "metano", "carbon footprint", "greenhouse gas"]},
    {"label": "Excreción de nitrógeno y fósforo", "terms": ["nitrogen excretion", "phosphorus excretion", "manure nutrients", "nutrient excretion"]},
    {"label": "Salud ruminal", "terms": ["rumen health", "salud ruminal", "ruminal fermentation", "rumen function"]},
    {"label": "Silajes y conservación de forrajes", "terms": ["silage", "ensilado", "forage preservation", "silage quality"]},
    {"label": "Digestibilidad y calidad del forraje", "terms": ["digestibility", "forage quality", "NDF", "digestibilidad"]},
    {"label": "Transición y posparto", "terms": ["transition cows", "posparto", "fresh cows", "transition period"]},
    {"label": "Cetosis y balance energético", "terms": ["ketosis", "negative energy balance", "BHB", "energy balance"]},
    {"label": "Acidosis", "terms": ["acidosis", "subacute ruminal acidosis", "SARA", "ruminal acidosis"]},
    {"label": "Mastitis y calidad de leche", "terms": ["mastitis", "milk quality", "somatic cell count", "udder health"]},
    {"label": "Reproducción y fertilidad", "terms": ["reproduction", "fertility", "reproductive performance", "estrus"]},
    {"label": "Longevidad y vida productiva", "terms": ["longevity", "productive life", "survival", "stayability"]},
    {"label": "Calostro e inmunidad neonatal", "terms": ["colostrum", "calostro", "passive immunity", "neonatal immunity"]},
    {"label": "Cría y recría", "terms": ["rearing", "recría", "young stock", "growing animals"]},
    {"label": "Destete y transición", "terms": ["weaning", "destete", "transition", "nursery"]},
    {"label": "Vitalidad neonatal y supervivencia", "terms": ["neonatal survival", "piglet vitality", "calf vitality", "early mortality"]},
    {"label": "Crecimiento y ganancia media diaria", "terms": ["growth", "average daily gain", "ADG", "performance"]},
    {"label": "Calidad de canal y de carne", "terms": ["carcass quality", "meat quality", "canal", "yield"]},
    {"label": "Calidad del huevo", "terms": ["egg quality", "shell quality", "albumen", "yolk"]},
    {"label": "Persistencia de puesta", "terms": ["laying persistence", "persistencia de puesta", "egg mass", "layers performance"]},
    {"label": "Eclosión e incubabilidad", "terms": ["hatchability", "incubability", "embryo development", "hatchery"]},
    {"label": "Coccidiosis", "terms": ["coccidiosis", "Eimeria", "anticoccidial", "intestinal lesions"]},
    {"label": "Salmonella", "terms": ["salmonella", "food safety", "preharvest control", "carrier state"]},
    {"label": "Influenza aviar", "terms": ["avian influenza", "HPAI", "influenza aviar", "avian flu"]},
    {"label": "Peste porcina africana", "terms": ["african swine fever", "ASF", "peste porcina africana", "ASFV"]},
    {"label": "PRRS y otras enfermedades porcinas", "terms": ["PRRS", "porcine reproductive and respiratory syndrome", "porcine diseases", "circovirus"]},
    {"label": "Lengua azul", "terms": ["bluetongue", "lengua azul", "vector-borne disease", "orbivirus"]},
    {"label": "Parasitismo en pequeños rumiantes", "terms": ["parasites", "parasitismo", "helminths", "small ruminants"]},
    {"label": "Enfermedades digestivas del conejo", "terms": ["rabbit enteropathy", "enteropatía", "rabbits", "digestive disorders"]},
    {"label": "Sensores y ganadería de precisión", "terms": ["precision livestock farming", "sensors", "monitoring", "digital phenotyping"]},
    {"label": "Modelos predictivos y datos", "terms": ["predictive models", "artificial intelligence", "machine learning", "decision support"]},
]

REG_TOPIC_ITEMS = [
    {"label": "Bienestar animal", "terms": ["bienestar animal", "animal welfare", "welfare rules", "condiciones de cría"]},
    {"label": "Bienestar en transporte", "terms": ["transporte animal", "animal transport", "transport welfare", "journey times"]},
    {"label": "Bienestar en sacrificio", "terms": ["sacrificio", "slaughter welfare", "stunning", "matadero"]},
    {"label": "Alojamiento y densidad", "terms": ["alojamiento", "stocking density", "housing requirements", "densidad"]},
    {"label": "Bioseguridad y sanidad animal", "terms": ["bioseguridad", "animal health", "farm hygiene", "sanidad animal"]},
    {"label": "Ley de Sanidad Animal de la UE", "terms": ["animal health law", "regulation 2016/429", "sanidad animal", "AHL"]},
    {"label": "Notificación obligatoria de enfermedades", "terms": ["notificación obligatoria", "notifiable diseases", "disease reporting", "declaración obligatoria"]},
    {"label": "Peste porcina africana", "terms": ["african swine fever", "ASF", "peste porcina africana", "regionalisation"]},
    {"label": "Influenza aviar", "terms": ["avian influenza", "influenza aviar", "HPAI", "avian flu"]},
    {"label": "Lengua azul", "terms": ["bluetongue", "lengua azul", "orbivirus", "movements"]},
    {"label": "Salmonella y zoonosis", "terms": ["salmonella", "zoonosis", "control programmes", "food safety"]},
    {"label": "Uso de antimicrobianos", "terms": ["antimicrobianos", "antimicrobials", "AMR", "prudent use"]},
    {"label": "Medicamentos veterinarios", "terms": ["veterinary medicines", "medicamentos veterinarios", "regulation 2019/6", "prescripción"]},
    {"label": "Piensos medicamentosos", "terms": ["medicated feed", "piensos medicamentosos", "regulation 2019/4", "carry-over"]},
    {"label": "Higiene de los piensos", "terms": ["feed hygiene", "higiene de los piensos", "regulation 183/2005", "HACCP"]},
    {"label": "Aditivos para alimentación animal", "terms": ["feed additives", "aditivos para piensos", "regulation 1831/2003", "authorisation"]},
    {"label": "Etiquetado de piensos", "terms": ["feed labelling", "etiquetado de piensos", "regulation 767/2009", "claims"]},
    {"label": "Materias primas para piensos", "terms": ["feed materials", "materias primas", "catalogue of feed materials", "raw materials"]},
    {"label": "Contaminantes y residuos", "terms": ["contaminants", "residues", "contaminantes", "residuos"]},
    {"label": "Micotoxinas y límites", "terms": ["mycotoxins", "micotoxinas", "guidance values", "DON"]},
    {"label": "Dioxinas y PCB", "terms": ["dioxins", "PCB", "contaminants", "feed safety"]},
    {"label": "Subproductos animales", "terms": ["animal by-products", "subproductos animales", "ABP", "Regulation 1069/2009"]},
    {"label": "Trazabilidad", "terms": ["trazabilidad", "traceability", "batch identification", "food chain"]},
    {"label": "Identificación y movimientos", "terms": ["movimientos", "animal identification", "traceability system", "movements"]},
    {"label": "Controles oficiales", "terms": ["official controls", "controles oficiales", "regulation 2017/625", "inspections"]},
    {"label": "Exportación y certificados sanitarios", "terms": ["export certificates", "certificados sanitarios", "third countries", "exports"]},
    {"label": "Importación y controles fronterizos", "terms": ["import controls", "border control posts", "BCP", "imports"]},
    {"label": "Producción ecológica", "terms": ["organic production", "producción ecológica", "organic livestock", "regulation 2018/848"]},
    {"label": "Ganadería y nitratos", "terms": ["nitrates", "nitratos", "water protection", "manure management"]},
    {"label": "Emisiones de amoniaco", "terms": ["ammonia emissions", "amoniaco", "air quality", "emissions"]},
    {"label": "Metano y gases de efecto invernadero", "terms": ["methane", "metano", "GHG", "greenhouse gases"]},
    {"label": "Autorizaciones ambientales", "terms": ["environmental permits", "autorización ambiental", "IPPC", "permits"]},
    {"label": "Emisiones industriales", "terms": ["industrial emissions", "IED", "emisiones industriales", "BAT"]},
    {"label": "Agua y vertidos", "terms": ["water", "discharges", "water framework", "vertidos"]},
    {"label": "Gestión de estiércoles", "terms": ["manure", "gestión de estiércoles", "slurry", "digestate"]},
    {"label": "Huella ambiental y reporting", "terms": ["environmental reporting", "sustainability reporting", "CSRD", "footprint"]},
    {"label": "Declaraciones y claims ambientales", "terms": ["environmental claims", "green claims", "sustainability claims", "claims"]},
    {"label": "Deforestación y materias primas importadas", "terms": ["deforestation", "EUDR", "soja", "imported raw materials"]},
    {"label": "Seguridad alimentaria pre y posgranja", "terms": ["food safety", "farm to fork", "preharvest", "food chain"]},
    {"label": "Huevos y normas de comercialización", "terms": ["egg marketing standards", "huevos", "egg labeling", "commercialisation"]},
    {"label": "Leche cruda y calidad higiénica", "terms": ["raw milk", "milk hygiene", "somatic cells", "quality standards"]},
    {"label": "Carne y requisitos de canal", "terms": ["meat hygiene", "carcass classification", "beef", "pork meat"]},
    {"label": "Bienestar en porcino", "terms": ["pig welfare", "tail docking", "porcino", "enrichment"]},
    {"label": "Bienestar en broilers", "terms": ["broiler welfare", "stocking density broilers", "pollos", "avicultura de carne"]},
    {"label": "Bienestar en ponedoras", "terms": ["laying hens welfare", "ponedoras", "cages", "egg production"]},
    {"label": "Bienestar en vacuno", "terms": ["cattle welfare", "bovine", "vacuno", "housing cattle"]},
    {"label": "Bienestar en pequeños rumiantes", "terms": ["sheep welfare", "goat welfare", "small ruminants", "ovino caprino"]},
    {"label": "Bienestar en conejos", "terms": ["rabbit welfare", "cunicultura", "rabbits", "housing rabbits"]},
    {"label": "Residuos de medicamentos", "terms": ["residues", "withdrawal period", "medicines residues", "LMR"]},
    {"label": "Resistencia antimicrobiana", "terms": ["antimicrobial resistance", "AMR", "resistencia antimicrobiana", "surveillance"]},
    {"label": "Normas de alimentación animal en España", "terms": ["alimentación animal", "piensos", "feed law", "normativa española"]},
    {"label": "Normativa de la EFSA y opiniones científicas", "terms": ["EFSA opinion", "scientific opinion", "risk assessment", "feed safety"]},
    {"label": "Ayudas PAC y ecoesquemas", "terms": ["PAC", "CAP", "ecoesquemas", "aids"]},
    {"label": "Condicionalidad y eco-regímenes", "terms": ["conditionality", "eco-schemes", "cross compliance", "greening"]},
    {"label": "Etiquetado de origen y comercialización", "terms": ["origin labelling", "country of origin", "commercialisation", "etiquetado"]},
    {"label": "Envases, residuos y sostenibilidad", "terms": ["packaging", "waste", "plastic", "circular economy"]},
]

MARKET_TOPICS = {item["label"]: item for item in MARKET_TOPIC_ITEMS}
SCIENCE_TOPICS = {item["label"]: item for item in SCIENCE_TOPIC_ITEMS}
REG_TOPICS = {item["label"]: item for item in REG_TOPIC_ITEMS}

DEFAULT_MARKET_SELECTION = [
    "Boletines oficiales de precios",
    "Coste del pienso",
    "Noticias del MAPA y del sector",
]
DEFAULT_SCIENCE_SELECTION = [
    "Nutrición y formulación",
    "Eficiencia alimentaria e índice de conversión",
    "Bioseguridad y prevención",
]
DEFAULT_REG_SELECTION = [
    "Bienestar animal",
    "Bioseguridad y sanidad animal",
    "Piensos, aditivos y alimentación animal",
] if "Piensos, aditivos y alimentación animal" in REG_TOPICS else [
    "Bienestar animal",
    "Bioseguridad y sanidad animal",
    "Aditivos para alimentación animal",
]

LEGAL_TERMS = {
    "reglamento", "directive", "directiva", "real decreto", "orden", "ley", "decisión", "decision", "boe",
    "eur-lex", "regulation", "implementing", "delegated", "resolución", "resolucion", "decreto", "ministerio",
}


def _strip_html(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", BeautifulSoup(text, "html.parser").get_text(" ", strip=True)).strip()



def _normalize_dt(value: Optional[datetime]) -> Optional[datetime]:
    if value is None:
        return None
    if value.tzinfo is not None:
        return value.astimezone(timezone.utc).replace(tzinfo=None)
    return value



def _parse_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    try:
        return _normalize_dt(date_parser.parse(str(value)))
    except Exception:
        return None



def _date_in_range(value: Optional[datetime], start: date, end: date) -> bool:
    if value is None:
        return True
    current = value.date()
    return start <= current <= end



def _truncate(text: str, max_len: int = 420) -> str:
    text = text or ""
    if len(text) <= max_len:
        return text
    return text[: max_len - 1].rstrip() + "…"



def _safe_get(url: str, **kwargs) -> requests.Response:
    headers = kwargs.pop("headers", {})
    merged = {"User-Agent": USER_AGENT}
    merged.update(headers)
    response = requests.get(url, timeout=REQUEST_TIMEOUT, headers=merged, **kwargs)
    response.raise_for_status()
    return response



def _clean_search_link(url: str) -> str:
    if not url:
        return ""
    if "duckduckgo.com/l/?" in url:
        parsed = urlparse(url)
        q = parse_qs(parsed.query)
        if "uddg" in q:
            return unquote(q["uddg"][0])
    return url



def _url_domain(url: str) -> str:
    try:
        return urlparse(url).netloc.lower()
    except Exception:
        return ""



def _keywords_from_text(text: str, top_k: int = 10) -> List[str]:
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚáéíóúÑñÜü0-9\-]{4,}", (text or "").lower())
    counts = Counter(token for token in tokens if token not in STOPWORDS)
    return [token for token, _ in counts.most_common(top_k)]



def _dedupe(records: List[dict]) -> List[dict]:
    seen = set()
    out: List[dict] = []
    for item in records:
        key = (
            (item.get("title") or "").strip().lower(),
            (item.get("url") or "").strip().lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out



def _all_species_selected(selected_species: Sequence[str]) -> bool:
    return not selected_species or "All species" in selected_species



def _selected_profiles(selected_species: Sequence[str]) -> List[Dict[str, List[str]]]:
    if _all_species_selected(selected_species):
        return [SPECIES_PROFILES["All species"], SPECIES_PROFILES["Alimentación animal"]]
    return [SPECIES_PROFILES[name] for name in selected_species if name in SPECIES_PROFILES]



def species_terms(selected_species: Sequence[str], category: str) -> List[str]:
    key = {
        "market": "market_aliases",
        "science": "science_aliases",
        "regulation": "reg_aliases",
    }[category]
    terms: List[str] = []
    for profile in _selected_profiles(selected_species):
        terms.extend(profile.get(key, []))
        terms.extend(profile.get("aliases", []))
    deduped = []
    seen = set()
    for term in terms:
        low = term.lower()
        if low not in seen:
            seen.add(low)
            deduped.append(term)
    return deduped[:18]



def species_filter_match(text: str, selected_species: Sequence[str], category: str) -> bool:
    if _all_species_selected(selected_species):
        return True
    hay = (text or "").lower()
    for term in species_terms(selected_species, category):
        if term.lower() in hay:
            return True
    if "Alimentación animal" in selected_species and any(token in hay for token in ["feed", "pienso", "alimentación animal", "animal nutrition"]):
        return True
    return False



def mentions_only_other_species(text: str, selected_species: Sequence[str], category: str) -> bool:
    if _all_species_selected(selected_species):
        return False
    hay = (text or "").lower()
    selected_aliases = {term.lower() for term in species_terms(selected_species, category)}
    if any(alias in hay for alias in selected_aliases):
        return False
    other_terms: List[str] = []
    for name, profile in SPECIES_PROFILES.items():
        if name in selected_species or name in {"All species", "Alimentación animal"}:
            continue
        other_terms.extend(profile.get("aliases", []))
    return any(term.lower() in hay for term in other_terms)



def topic_terms(topic_dict: Dict[str, dict], selected_labels: Sequence[str]) -> List[str]:
    terms: List[str] = []
    for label in selected_labels:
        item = topic_dict.get(label)
        if item:
            terms.extend(item.get("terms", []))
            terms.append(label)
    deduped = []
    seen = set()
    for term in terms:
        low = term.lower()
        if low not in seen:
            seen.add(low)
            deduped.append(term)
    return deduped



def _topic_hit_count(text: str, terms: Sequence[str]) -> int:
    hay = (text or "").lower()
    return sum(1 for term in terms if term.lower() in hay)



def _official_domain(url: str) -> bool:
    domain = _url_domain(url)
    return any(domain.endswith(item) for item in OFFICIAL_REG_DOMAINS)


@st.cache_data(show_spinner=False, ttl=7200)
def search_google_news(query: str, start_date: date, end_date: date, max_results: int = 12) -> List[dict]:
    url = f"https://news.google.com/rss/search?q={quote(query)}&hl=es&gl=ES&ceid=ES:es"
    response = _safe_get(url)
    feed = feedparser.parse(response.content)
    records: List[dict] = []
    for entry in getattr(feed, "entries", []):
        published = _parse_date(entry.get("published") or entry.get("pubDate"))
        if not _date_in_range(published, start_date, end_date):
            continue
        source = "Google News"
        source_obj = entry.get("source")
        if isinstance(source_obj, dict):
            source = source_obj.get("title", source)
        records.append(
            {
                "title": _strip_html(entry.get("title", "Sin título")),
                "snippet": _truncate(_strip_html(entry.get("summary", ""))),
                "url": entry.get("link", ""),
                "source": source,
                "published": published.isoformat() if published else "",
                "query": query,
                "search_engine": "Google News",
            }
        )
        if len(records) >= max_results:
            break
    return _dedupe(records)


@st.cache_data(show_spinner=False, ttl=7200)
def search_duckduckgo_web(query: str, max_results: int = 10) -> List[dict]:
    response = _safe_get("https://duckduckgo.com/html/", params={"q": query, "kl": "es-es"})
    soup = BeautifulSoup(response.text, "html.parser")
    records: List[dict] = []
    for result in soup.select("div.result"):
        a = result.select_one("a.result__a")
        if not a:
            continue
        raw_url = a.get("href", "")
        url = _clean_search_link(raw_url)
        title = _strip_html(a.get_text(" ", strip=True))
        snippet_node = result.select_one("a.result__snippet") or result.select_one("div.result__snippet")
        snippet = _truncate(_strip_html(snippet_node.get_text(" ", strip=True)) if snippet_node else "")
        records.append(
            {
                "title": title,
                "snippet": snippet,
                "url": url,
                "source": _url_domain(url) or "DuckDuckGo",
                "published": "",
                "query": query,
                "search_engine": "DuckDuckGo",
            }
        )
        if len(records) >= max_results:
            break
    return _dedupe(records)


@st.cache_data(show_spinner=False, ttl=7200)
def search_openalex(query: str, start_date: date, end_date: date, max_results: int = 10) -> List[dict]:
    params = {
        "search": query,
        "filter": f"from_publication_date:{start_date.isoformat()},to_publication_date:{end_date.isoformat()}",
        "per-page": max_results,
        "sort": "relevance_score:desc",
        "mailto": "radar@example.com",
    }
    response = _safe_get("https://api.openalex.org/works", params=params)
    data = response.json()
    records: List[dict] = []
    for item in data.get("results", []):
        published = item.get("publication_date") or str(item.get("publication_year") or "")
        abstract = ""
        inverted = item.get("abstract_inverted_index") or {}
        if inverted:
            positions = []
            for word, indices in inverted.items():
                for idx in indices:
                    positions.append((idx, word))
            abstract = " ".join(word for _, word in sorted(positions)[:120])
        source = "OpenAlex"
        primary = item.get("primary_location") or {}
        if primary.get("source") and primary["source"].get("display_name"):
            source = primary["source"]["display_name"]
        authors = ", ".join(a.get("author", {}).get("display_name", "") for a in item.get("authorships", [])[:6]).strip(", ")
        records.append(
            {
                "title": item.get("display_name", "Sin título"),
                "snippet": _truncate(abstract or item.get("title", "")),
                "url": item.get("doi") or (primary.get("landing_page_url") or item.get("id", "")),
                "source": source,
                "published": published,
                "authors": authors,
                "journal": source,
                "doi": item.get("doi", ""),
                "query": query,
                "search_engine": "OpenAlex",
            }
        )
    return _dedupe(records)


@st.cache_data(show_spinner=False, ttl=7200)
def search_europe_pmc(query: str, start_date: date, end_date: date, max_results: int = 10) -> List[dict]:
    full_query = f"({query}) AND FIRST_PDATE:[{start_date.isoformat()} TO {end_date.isoformat()}]"
    params = {"query": full_query, "format": "json", "pageSize": max_results, "sort": "FIRST_PDATE_D"}
    response = _safe_get("https://www.ebi.ac.uk/europepmc/webservices/rest/search", params=params)
    data = response.json()
    records: List[dict] = []
    for item in data.get("resultList", {}).get("result", []):
        published = item.get("firstPublicationDate") or item.get("pubYear", "")
        doi = item.get("doi")
        url = ""
        if doi:
            url = f"https://doi.org/{doi}"
        elif item.get("pmid"):
            url = f"https://europepmc.org/article/MED/{item['pmid']}"
        elif item.get("pmcid"):
            url = f"https://europepmc.org/article/PMC/{item['pmcid']}"
        records.append(
            {
                "title": _strip_html(item.get("title", "Sin título")),
                "snippet": _truncate(_strip_html(item.get("abstractText", "")) or item.get("journalTitle", "")),
                "url": url,
                "source": item.get("journalTitle", "Europe PMC"),
                "published": published,
                "authors": item.get("authorString", ""),
                "journal": item.get("journalTitle", "Europe PMC"),
                "doi": doi or "",
                "query": query,
                "search_engine": "Europe PMC",
            }
        )
    return _dedupe(records)



def build_queries(selected_species: Sequence[str], selected_topics: Sequence[str], category: str) -> List[str]:
    species_clause = species_terms(selected_species, category)
    species_phrase = " OR ".join(f'"{term}"' for term in species_clause[:8]) if species_clause else '"ganadería"'
    topic_catalog = {
        "market": MARKET_TOPICS,
        "science": SCIENCE_TOPICS,
        "regulation": REG_TOPICS,
    }[category]
    queries: List[str] = []
    for label in selected_topics:
        item = topic_catalog.get(label)
        if not item:
            continue
        topic_phrase = " OR ".join(f'"{term}"' for term in item["terms"][:6])
        if category == "market":
            queries.append(f"({species_phrase}) ({topic_phrase})")
            queries.append(f"({species_phrase}) ({topic_phrase}) (MAPA OR ministerio agricultura OR boletín OR precios OR mercado OR noticias)")
        elif category == "science":
            species_simple = " ".join(species_clause[:4]) if species_clause else "livestock animal nutrition"
            topic_simple = " ".join(item["terms"][:5])
            queries.append(f"{species_simple} {topic_simple}")
        else:
            official_sites = "(site:boe.es OR site:eur-lex.europa.eu OR site:mapa.gob.es OR site:aesan.gob.es OR site:efsa.europa.eu OR site:miteco.gob.es OR site:sanidad.gob.es)"
            legal_anchor = "(reglamento OR directiva OR ley OR real decreto OR orden OR resolución OR BOE OR EUR-Lex OR regulation)"
            queries.append(f"({species_phrase}) ({topic_phrase}) {legal_anchor} {official_sites}")
            queries.append(f"({species_phrase}) ({topic_phrase}) normativa legislación regulación {official_sites}")
    deduped = []
    seen = set()
    for query in queries:
        low = query.lower().strip()
        if low and low not in seen:
            seen.add(low)
            deduped.append(query)
    return deduped[:24]



def score_result(item: dict, category: str, selected_species: Sequence[str], selected_topic_labels: Sequence[str]) -> int:
    text = " ".join(
        [
            item.get("title", ""),
            item.get("snippet", ""),
            item.get("source", ""),
            item.get("url", ""),
            item.get("query", ""),
        ]
    ).lower()
    if category == "market":
        topic_dict = MARKET_TOPICS
    elif category == "science":
        topic_dict = SCIENCE_TOPICS
    else:
        topic_dict = REG_TOPICS
    terms = topic_terms(topic_dict, selected_topic_labels)
    topic_hits = _topic_hit_count(text, terms)
    species_ok = species_filter_match(text, selected_species, category)
    score = topic_hits * 3
    if species_ok:
        score += 4
    if mentions_only_other_species(text, selected_species, category):
        score -= 8
    domain = _url_domain(item.get("url", ""))
    if category == "market":
        if any(domain.endswith(site) for site in MARKET_NEWS_SITES):
            score += 3
        if any(token in text for token in ["precio", "prices", "market", "mercado", "boletín", "cotización", "news"]):
            score += 2
    elif category == "science":
        if item.get("search_engine") in {"OpenAlex", "Europe PMC"}:
            score += 6
        if any(hint in text for hint in SCIENCE_HINTS):
            score += 2
    else:
        if _official_domain(item.get("url", "")):
            score += 8
        legal_hits = sum(1 for token in LEGAL_TERMS if token in text)
        score += legal_hits * 2
        if legal_hits == 0:
            score -= 6
    return score



def filter_and_rank(records: List[dict], category: str, selected_species: Sequence[str], selected_topics: Sequence[str], max_results: int) -> List[dict]:
    ranked: List[tuple] = []
    for item in _dedupe(records):
        text = " ".join([item.get("title", ""), item.get("snippet", ""), item.get("source", ""), item.get("url", "")])
        score = score_result(item, category, selected_species, selected_topics)
        topic_catalog = {"market": MARKET_TOPICS, "science": SCIENCE_TOPICS, "regulation": REG_TOPICS}[category]
        term_list = topic_terms(topic_catalog, selected_topics)
        topic_hit = _topic_hit_count(text, term_list) > 0
        species_ok = species_filter_match(text, selected_species, category)

        if category == "market":
            if not topic_hit:
                continue
            if not _all_species_selected(selected_species) and not species_ok and "Alimentación animal" not in selected_species:
                continue
            if score < 4:
                continue
        elif category == "science":
            if score < 5:
                continue
        else:
            if not _official_domain(item.get("url", "")):
                continue
            if not topic_hit:
                continue
            if not _all_species_selected(selected_species) and not species_ok and "Alimentación animal" not in selected_species:
                continue
            if mentions_only_other_species(text, selected_species, category):
                continue
            if score < 10:
                continue

        enriched = dict(item)
        enriched["score"] = score
        ranked.append((score, enriched))

    ranked.sort(key=lambda x: (x[0], x[1].get("published", "")), reverse=True)
    return [item for _, item in ranked[:max_results]]



def search_market(selected_species: Sequence[str], selected_topics: Sequence[str], start_date: date, end_date: date, max_results: int) -> List[dict]:
    queries = build_queries(selected_species, selected_topics, "market")
    records: List[dict] = []
    for query in queries:
        try:
            records.extend(search_google_news(query, start_date, end_date, max_results=6))
        except Exception:
            pass
        try:
            records.extend(search_duckduckgo_web(query, max_results=5))
        except Exception:
            pass
    return filter_and_rank(records, "market", selected_species, selected_topics, max_results)



def search_science(selected_species: Sequence[str], selected_topics: Sequence[str], start_date: date, end_date: date, max_results: int) -> List[dict]:
    queries = build_queries(selected_species, selected_topics, "science")
    records: List[dict] = []
    for query in queries:
        try:
            records.extend(search_openalex(query, start_date, end_date, max_results=6))
        except Exception:
            pass
        try:
            records.extend(search_europe_pmc(query, start_date, end_date, max_results=6))
        except Exception:
            pass
    return filter_and_rank(records, "science", selected_species, selected_topics, max_results)



def search_regulation(selected_species: Sequence[str], selected_topics: Sequence[str], start_date: date, end_date: date, max_results: int) -> List[dict]:
    queries = build_queries(selected_species, selected_topics, "regulation")
    records: List[dict] = []
    for query in queries:
        try:
            records.extend(search_duckduckgo_web(query, max_results=8))
        except Exception:
            pass
        try:
            records.extend(search_google_news(query, start_date, end_date, max_results=5))
        except Exception:
            pass
    return filter_and_rank(records, "regulation", selected_species, selected_topics, max_results)



def run_search(
    selected_species: Sequence[str],
    market_topics: Sequence[str],
    science_topics: Sequence[str],
    reg_topics: Sequence[str],
    start_date: date,
    end_date: date,
    max_results: int,
    only_category: Optional[str] = None,
    existing: Optional[Dict[str, List[dict]]] = None,
) -> Dict[str, List[dict]]:
    results = existing or {"market": [], "science": [], "regulation": []}
    if only_category in (None, "market"):
        results["market"] = search_market(selected_species, market_topics, start_date, end_date, max_results) if market_topics else []
    if only_category in (None, "science"):
        results["science"] = search_science(selected_species, science_topics, start_date, end_date, max_results) if science_topics else []
    if only_category in (None, "regulation"):
        results["regulation"] = search_regulation(selected_species, reg_topics, start_date, end_date, max_results) if reg_topics else []
    return results



def flatten_results(results: Dict[str, List[dict]]) -> List[dict]:
    rows: List[dict] = []
    for category, items in results.items():
        for item in items:
            enriched = dict(item)
            enriched["category"] = CATEGORY_LABELS[category]
            rows.append(enriched)
    return rows



def corpus_text(results: Dict[str, List[dict]], limit_per_category: int = 8) -> str:
    lines = []
    for category in ["market", "science", "regulation"]:
        lines.append(f"\n## {CATEGORY_LABELS[category]}\n")
        for idx, item in enumerate(results.get(category, [])[:limit_per_category], start=1):
            lines.append(
                f"[{idx}] {item.get('title', '')}\n"
                f"Fuente: {item.get('source', '')}\n"
                f"Fecha: {item.get('published', '')[:10]}\n"
                f"Resumen: {item.get('snippet', '')}\n"
                f"URL: {item.get('url', '')}\n"
            )
    return "\n".join(lines)



def llm_is_available() -> bool:
    return bool(os.getenv("OPENAI_API_KEY")) and OpenAI is not None



def call_openai(system_prompt: str, user_prompt: str) -> str:
    if not llm_is_available():
        raise RuntimeError("No hay configuración de OpenAI disponible.")
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    model = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
    try:
        response = client.responses.create(model=model, instructions=system_prompt, input=user_prompt)
        if getattr(response, "output_text", None):
            return response.output_text.strip()
    except Exception:
        pass
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    return (response.choices[0].message.content or "").strip()



def extractive_brief(
    selected_species: Sequence[str],
    selected_market_topics: Sequence[str],
    selected_science_topics: Sequence[str],
    selected_reg_topics: Sequence[str],
    results: Dict[str, List[dict]],
    company_context: str,
) -> str:
    flat = flatten_results(results)
    if not flat:
        return "No hay resultados suficientes para elaborar el briefing."

    corpus = " ".join((item.get("title", "") + " " + item.get("snippet", "")) for item in flat)
    hot_terms = _keywords_from_text(corpus, top_k=12)
    species_text = ", ".join(selected_species) if selected_species else "All species"

    lines = [
        f"# Briefing radar | {species_text}",
        "",
        "## Resumen ejecutivo",
        f"Se han revisado **{len(results.get('market', []))}** señales de mercado/noticias, **{len(results.get('science', []))}** referencias científico-técnicas y **{len(results.get('regulation', []))}** referencias regulatorias.",
        f"Especies/segmentos analizados: **{species_text}**.",
        f"Temas seleccionados en mercado: {', '.join(selected_market_topics[:6]) if selected_market_topics else 'sin selección'}.",
        f"Temas seleccionados en científico-técnico: {', '.join(selected_science_topics[:6]) if selected_science_topics else 'sin selección'}.",
        f"Temas seleccionados en regulación: {', '.join(selected_reg_topics[:6]) if selected_reg_topics else 'sin selección'}.",
        f"Señales repetidas en las fuentes recuperadas: {', '.join(hot_terms[:8]) if hot_terms else 'sin patrón claro'}.",
        "",
        "## Lectura integrada para marketing",
        "El radar debe leerse como una combinación de presión de mercado, evidencia técnica y restricciones regulatorias para traducir necesidades externas en productos, soluciones y argumentarios de Nutreco Iberia.",
    ]

    if results.get("market"):
        lines.extend(["", "## Mercado y noticias"]) 
        for item in results["market"][:5]:
            lines.append(f"- **{item['title']}** ({item.get('source', 'Fuente')}). {item.get('snippet', '')}")
    else:
        lines.extend(["", "## Mercado y noticias", "- No se han recuperado resultados suficientemente relevantes con la selección actual."])

    if results.get("science"):
        lines.extend(["", "## Evidencia científico-técnica"]) 
        for item in results["science"][:5]:
            source = item.get("journal") or item.get("source", "Fuente científica")
            lines.append(f"- **{item['title']}** ({source}). {item.get('snippet', '')}")
    else:
        lines.extend(["", "## Evidencia científico-técnica", "- No se han recuperado publicaciones suficientemente relevantes con la selección actual."])

    if results.get("regulation"):
        lines.extend(["", "## Legislación y regulación"]) 
        for item in results["regulation"][:5]:
            lines.append(f"- **{item['title']}** ({item.get('source', 'Fuente oficial')}). {item.get('snippet', '')}")
    else:
        lines.extend(["", "## Legislación y regulación", "- No se han recuperado referencias regulatorias suficientemente relevantes con la selección actual."])

    lines.extend(
        [
            "",
            "## Implicaciones preliminares para Nutreco Iberia",
            "- Revisar si las señales de mercado justifican mensajes, soluciones o argumentarios específicos por especie o de forma transversal en alimentación animal.",
            "- Contrastar las oportunidades de producto con la evidencia científico-técnica más sólida antes de priorizar un claim o una propuesta de valor.",
            "- Verificar la viabilidad regulatoria de cada solución potencial antes de cualquier despliegue comercial.",
            "- Priorizar vigilancia continua en los temas con presencia simultánea en mercado, ciencia y regulación.",
            "",
            "## Contexto corporativo utilizado",
            company_context.strip(),
        ]
    )
    return "\n".join(lines)



def generate_brief(
    selected_species: Sequence[str],
    selected_market_topics: Sequence[str],
    selected_science_topics: Sequence[str],
    selected_reg_topics: Sequence[str],
    results: Dict[str, List[dict]],
    company_context: str,
) -> str:
    if not llm_is_available():
        return extractive_brief(
            selected_species,
            selected_market_topics,
            selected_science_topics,
            selected_reg_topics,
            results,
            company_context,
        )

    system_prompt = (
        "Eres un analista senior de inteligencia de mercado, ciencia aplicada y regulación para nutrición animal. "
        "Usa solo el corpus suministrado. No inventes datos. Escribe en español con tono ejecutivo y práctico."
    )
    user_prompt = f"""
Especies/segmentos: {', '.join(selected_species) if selected_species else 'All species'}
Temas de mercado: {', '.join(selected_market_topics)}
Temas científico-técnicos: {', '.join(selected_science_topics)}
Temas regulatorios: {', '.join(selected_reg_topics)}

Contexto corporativo:
{company_context}

Corpus:
{corpus_text(results)}

Estructura obligatoria:
1. Resumen ejecutivo.
2. Lectura integrada para marketing.
3. Señales de mercado con implicaciones de portafolio.
4. Hallazgos científico-técnicos con posible traducción a soluciones.
5. Implicaciones regulatorias.
6. Recomendaciones priorizadas para Nutreco Iberia.
7. Riesgos, vacíos y siguientes preguntas.
"""
    return call_openai(system_prompt, user_prompt)



def bibliography_entries(results: Dict[str, List[dict]]) -> List[str]:
    entries: List[str] = []
    for item in flatten_results(results):
        published = (item.get("published") or "")[:10] or "s/f"
        if item.get("category") == CATEGORY_LABELS["science"]:
            authors = item.get("authors") or "Autoría no disponible"
            journal = item.get("journal") or item.get("source") or "Fuente científica"
            doi_or_url = item.get("doi") or item.get("url", "")
            entries.append(f"{authors}. ({published}). {item.get('title')}. {journal}. {doi_or_url}")
        else:
            entries.append(f"{item.get('source', 'Fuente no indicada')}. ({published}). {item.get('title')}. {item.get('url', '')}")
    return entries



def build_docx_bytes(
    selected_species: Sequence[str],
    selected_market_topics: Sequence[str],
    selected_science_topics: Sequence[str],
    selected_reg_topics: Sequence[str],
    start_date: date,
    end_date: date,
    company_context: str,
    brief_text: str,
    results: Dict[str, List[dict]],
) -> bytes:
    doc = Document()
    species_text = ", ".join(selected_species) if selected_species else "All species"
    doc.add_heading(f"Radar sectorial | {species_text}", level=0)
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Intervalo analizado: {start_date.isoformat()} a {end_date.isoformat()}")
    doc.add_paragraph(f"Mercado/noticias: {', '.join(selected_market_topics) if selected_market_topics else 'Sin selección'}")
    doc.add_paragraph(f"Científico-técnico: {', '.join(selected_science_topics) if selected_science_topics else 'Sin selección'}")
    doc.add_paragraph(f"Regulación: {', '.join(selected_reg_topics) if selected_reg_topics else 'Sin selección'}")

    doc.add_heading("Briefing", level=1)
    for paragraph in brief_text.split("\n"):
        clean = paragraph.strip()
        if not clean:
            doc.add_paragraph("")
            continue
        if clean.startswith("# "):
            doc.add_heading(clean[2:], level=1)
        elif clean.startswith("## "):
            doc.add_heading(clean[3:], level=2)
        elif clean.startswith("- "):
            doc.add_paragraph(clean[2:], style="List Bullet")
        else:
            doc.add_paragraph(clean)

    doc.add_heading("Contexto corporativo aplicado", level=1)
    doc.add_paragraph(company_context.strip())

    doc.add_heading("Fuentes recuperadas", level=1)
    for category in ["market", "science", "regulation"]:
        doc.add_heading(CATEGORY_LABELS[category], level=2)
        items = results.get(category, [])
        if not items:
            doc.add_paragraph("Sin resultados.")
            continue
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item.get("title", "Sin título")).bold = True
            p.add_run(f" | {item.get('source', 'Fuente')} | {str(item.get('published', ''))[:10]}\n")
            p.add_run(item.get("snippet", ""))
            if item.get("url"):
                p.add_run(f"\n{item['url']}")

    doc.add_heading("Referencias bibliográficas / documentales", level=1)
    for entry in bibliography_entries(results):
        doc.add_paragraph(entry, style="List Number")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



def render_results(items: List[dict], empty_text: str) -> None:
    if not items:
        st.info(empty_text)
        return
    for item in items:
        title = item.get("title", "Sin título")
        url = item.get("url", "")
        source = item.get("source", "Fuente")
        published = (item.get("published") or "")[:10] or "s/f"
        score = item.get("score", "")
        with st.container():
            if url:
                st.markdown(f"**[{title}]({url})**")
            else:
                st.markdown(f"**{title}**")
            st.caption(f"{source} | {published} | Relevancia interna: {score}")
            st.write(item.get("snippet", ""))
            with st.expander("Ver metadatos"):
                st.write({
                    "query": item.get("query", ""),
                    "motor": item.get("search_engine", ""),
                    "url": url,
                })
            st.markdown("---")



def load_readme_text() -> str:
    readme_path = os.path.join(os.path.dirname(__file__), "README.md")
    if os.path.exists(readme_path):
        with open(readme_path, "r", encoding="utf-8") as handle:
            return handle.read()
    return "README no disponible."



def init_state() -> None:
    st.session_state.setdefault("search_results", {"market": [], "science": [], "regulation": []})
    st.session_state.setdefault("brief_text", "")
    st.session_state.setdefault("last_filters", {})



def reset_state() -> None:
    st.session_state["search_results"] = {"market": [], "science": [], "regulation": []}
    st.session_state["brief_text"] = ""
    st.session_state["last_filters"] = {}



def summary_metrics(results: Dict[str, List[dict]]) -> pd.DataFrame:
    rows = []
    for category in ["market", "science", "regulation"]:
        items = results.get(category, [])
        rows.append(
            {
                "Bloque": CATEGORY_LABELS[category],
                "Resultados": len(items),
                "Fuentes distintas": len({item.get('source', '') for item in items}),
            }
        )
    return pd.DataFrame(rows)



def main() -> None:
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    init_state()

    st.title(APP_TITLE)
    st.caption(
        "Radar para el departamento de marketing: convierte señales de mercado, evidencia técnica y novedades regulatorias en oportunidades de producto, soluciones y argumentario para Nutreco Iberia."
    )

    with st.sidebar:
        st.header("1. Alcance del radar")
        species_selection = st.multiselect(
            "Especies / segmentos",
            options=list(SPECIES_PROFILES.keys()),
            default=["Alimentación animal"],
            help="Puedes seleccionar una o varias especies. Si eliges 'All species', la búsqueda será transversal.",
        )
        today = date.today()
        default_start = date(today.year, max(1, today.month - 3), 1)
        start_date = st.date_input("Fecha inicio", value=default_start)
        end_date = st.date_input("Fecha fin", value=today)
        max_results = st.slider("Máximo de resultados por bloque", min_value=8, max_value=30, value=16, step=1)
        company_context = st.text_area(
            "Contexto corporativo / criterio de lectura",
            value=DEFAULT_COMPANY_CONTEXT,
            height=180,
        )
        col_a, col_b = st.columns(2)
        search_all = col_a.button("Buscar todo el radar", use_container_width=True)
        reset_button = col_b.button("Empezar de nuevo", use_container_width=True)

    if reset_button:
        reset_state()
        st.success("Se ha limpiado la búsqueda anterior.")

    if start_date > end_date:
        st.error("La fecha inicial no puede ser posterior a la fecha final.")
        return

    st.subheader("2. Propuestas de búsqueda por bloque")
    st.write(
        "Selecciona propuestas curadas. Cada selección lanza búsquedas más precisas que el antiguo campo libre y reduce ruido por especie y por tema."
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        market_topics = st.multiselect(
            f"Mercado y noticias ({len(MARKET_TOPICS)} opciones)",
            options=list(MARKET_TOPICS.keys()),
            default=DEFAULT_MARKET_SELECTION,
            help="Selecciona una o varias necesidades de vigilancia de mercado.",
        )
        search_market_button = st.button("Buscar mercado / noticias", use_container_width=True)
    with col2:
        science_topics = st.multiselect(
            f"Científico-técnico ({len(SCIENCE_TOPICS)} opciones)",
            options=list(SCIENCE_TOPICS.keys()),
            default=DEFAULT_SCIENCE_SELECTION,
            help="Selecciona una o varias líneas científicas.",
        )
        search_science_button = st.button("Buscar científico-técnico", use_container_width=True)
    with col3:
        reg_topics = st.multiselect(
            f"Legislación y regulación ({len(REG_TOPICS)} opciones)",
            options=list(REG_TOPICS.keys()),
            default=DEFAULT_REG_SELECTION,
            help="Selecciona uno o varios frentes regulatorios.",
        )
        search_reg_button = st.button("Buscar regulación", use_container_width=True)

    if not species_selection:
        species_selection = ["All species"]

    if search_all:
        with st.spinner("Recuperando mercado, ciencia y regulación..."):
            try:
                st.session_state.search_results = run_search(
                    species_selection,
                    market_topics,
                    science_topics,
                    reg_topics,
                    start_date,
                    end_date,
                    max_results,
                    only_category=None,
                    existing={"market": [], "science": [], "regulation": []},
                )
                st.session_state.brief_text = ""
                st.session_state.last_filters = {
                    "species": species_selection,
                    "market_topics": market_topics,
                    "science_topics": science_topics,
                    "reg_topics": reg_topics,
                    "start_date": start_date,
                    "end_date": end_date,
                    "company_context": company_context,
                }
                st.success("Radar actualizado en los tres bloques.")
            except Exception as exc:
                st.error(f"No se pudo completar la búsqueda: {exc}")

    if search_market_button:
        with st.spinner("Buscando mercado y noticias..."):
            try:
                st.session_state.search_results = run_search(
                    species_selection,
                    market_topics,
                    science_topics,
                    reg_topics,
                    start_date,
                    end_date,
                    max_results,
                    only_category="market",
                    existing=st.session_state.search_results,
                )
                st.session_state.brief_text = ""
                st.success("Bloque de mercado actualizado.")
            except Exception as exc:
                st.error(f"No se pudo completar la búsqueda de mercado: {exc}")

    if search_science_button:
        with st.spinner("Buscando evidencia científico-técnica..."):
            try:
                st.session_state.search_results = run_search(
                    species_selection,
                    market_topics,
                    science_topics,
                    reg_topics,
                    start_date,
                    end_date,
                    max_results,
                    only_category="science",
                    existing=st.session_state.search_results,
                )
                st.session_state.brief_text = ""
                st.success("Bloque científico-técnico actualizado.")
            except Exception as exc:
                st.error(f"No se pudo completar la búsqueda científica: {exc}")

    if search_reg_button:
        with st.spinner("Buscando normativa y regulación..."):
            try:
                st.session_state.search_results = run_search(
                    species_selection,
                    market_topics,
                    science_topics,
                    reg_topics,
                    start_date,
                    end_date,
                    max_results,
                    only_category="regulation",
                    existing=st.session_state.search_results,
                )
                st.session_state.brief_text = ""
                st.success("Bloque regulatorio actualizado.")
            except Exception as exc:
                st.error(f"No se pudo completar la búsqueda regulatoria: {exc}")

    results = st.session_state.search_results
    st.subheader("3. Resultado del radar")
    metrics_df = summary_metrics(results)
    st.dataframe(metrics_df, use_container_width=True, hide_index=True)

    tabs = st.tabs([
        "Mercado y noticias",
        "Científico-técnico",
        "Legislación y regulación",
        "Briefing e informe",
        "README / ayuda",
    ])

    with tabs[0]:
        st.caption("Prensa sectorial, boletines, noticias de mercado y actualizaciones institucionales relacionadas con la selección activa.")
        render_results(results.get("market", []), "No hay resultados de mercado relevantes con la selección actual.")

    with tabs[1]:
        st.caption("Publicaciones científicas y técnicas priorizadas por especie/segmento y por propuesta curada de búsqueda.")
        render_results(results.get("science", []), "No hay resultados científico-técnicos relevantes con la selección actual.")

    with tabs[2]:
        st.caption("Resultados regulatorios filtrados para evitar ruido de otras especies y priorizar dominios oficiales.")
        render_results(results.get("regulation", []), "No hay resultados regulatorios relevantes con la selección actual.")

    with tabs[3]:
        col_l, col_r = st.columns([1, 1])
        if col_l.button("Generar briefing", use_container_width=True):
            try:
                st.session_state.brief_text = generate_brief(
                    species_selection,
                    market_topics,
                    science_topics,
                    reg_topics,
                    results,
                    company_context,
                )
            except Exception as exc:
                st.error(f"No se pudo generar el briefing: {exc}")

        if st.session_state.brief_text:
            st.markdown(st.session_state.brief_text)
            docx_bytes = build_docx_bytes(
                species_selection,
                market_topics,
                science_topics,
                reg_topics,
                start_date,
                end_date,
                company_context,
                st.session_state.brief_text,
                results,
            )
            col_r.download_button(
                label="Descargar informe Word (.docx)",
                data=docx_bytes,
                file_name="radar_nutreco_iberia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            with st.expander("Referencias bibliográficas y documentales"):
                for entry in bibliography_entries(results):
                    st.markdown(f"- {entry}")
        else:
            st.info("Pulsa 'Generar briefing' cuando tengas resultados suficientes.")

    with tabs[4]:
        st.markdown(load_readme_text())

    st.divider()
    st.caption(
        "Aviso: el radar es una herramienta de vigilancia y priorización. Antes de usar cualquier conclusión en materiales externos, claims o recomendaciones comerciales, valida cada punto con la fuente primaria."
    )


if __name__ == "__main__":
    main()
